import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Set standard 1 inch margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Helper function to style table headers
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

# Title
title_p = doc.add_paragraph()
title_run = title_p.add_run("MODO RMC BATCH SHEET RECONCILIATION & MATERIAL CONSUMPTION ENGINE")
title_run.font.name = "Arial"
title_run.font.size = Pt(20)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(30, 58, 138) # Dark Navy
title_p.paragraph_format.space_after = Pt(4)

subtitle_p = doc.add_paragraph()
sub_run = subtitle_p.add_run("Architecture, Plant Drivers & Material Consumption Automation Guide")
sub_run.font.name = "Arial"
sub_run.font.size = Pt(13)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate
subtitle_p.paragraph_format.space_after = Pt(18)

# Section 1: Executive Summary
h1 = doc.add_heading("1. Executive Summary & Objective", level=1)
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph(
    "Ready-Mix Concrete (RMC) operations produce physical and digital autographic batch sheets from disparate batching control systems (e.g. Schwing Stetter EWIG/M1.5, MCI 70, MCI 360). "
    "The primary challenge has been the tedious manual entry of batch weights and material consumptions into ERP inventory tables. "
    "We have engineered an intelligent, enterprise-grade Batch Sheet Reconciliation Engine that delivers zero-developer plant onboarding, real-time split-screen verification, and automated material consumption reconciliation."
)
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(10)

# Section 2: Core Focus - Material Consumption Reconciliation
h1 = doc.add_heading("2. Core Focus: Material Consumption Reconciliation", level=1)
h1.paragraph_format.space_before = Pt(14)
h1.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph(
    "The engine specifically solves the inventory consumption reconciliation bottleneck using the 'Map Once, Reconcile Automatically Forever' paradigm:"
)
p.paragraph_format.space_after = Pt(6)

bp1 = doc.add_paragraph(style='List Bullet')
r = bp1.add_run("Map Once per Plant: ")
r.bold = True
bp1.add_run("When an operator first uploads a batch sheet from a given plant, they map each sheet label (e.g. MSA1, 12MM, CEM Silo 1, WAT, ADM 1, Silica) to the corresponding ERP Master Product dropdown.")

bp2 = doc.add_paragraph(style='List Bullet')
r = bp2.add_run("Permanent Memory: ")
r.bold = True
bp2.add_run("Clicking 'Remember Mappings for this Plant' permanently stores the column-to-product mapping in mm_batch_sheet_templates.material_mapping.")

bp3 = doc.add_paragraph(style='List Bullet')
r = bp3.add_run("Zero-Click Automation on Future Uploads: ")
r.bold = True
bp3.add_run("Every subsequent batch sheet uploaded from that plant automatically resolves 100% of material lines, compares actual vs target weights, checks tolerance, and prepares the consumption entry with zero manual user interaction.")

bp4 = doc.add_paragraph(style='List Bullet')
r = bp4.add_run("Atomic Batch Posting: ")
r.bold = True
bp4.add_run("Upon clicking 'Approve & Save Batch', the system posts directly to mm_batches and mm_batch_materials, updating actual inventory deductions.")

# Section 3: Scalable Plant Driver Architecture (1,000+ Plants)
h1 = doc.add_heading("3. Scalable Plant Driver Architecture (1,000+ Plants)", level=1)
h1.paragraph_format.space_before = Pt(14)
h1.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph(
    "To support 1,000+ distinct plants without central bottlenecks or database seeders, the system uses dynamic file auto-discovery:"
)

# Table of Drivers
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

hdr_cells = table.rows[0].cells
headers = ["Driver File Name", "Plant / Brand Target", "Mixer Size", "Extraction Pattern"]
widths = [Inches(2.2), Inches(1.8), Inches(1.0), Inches(1.5)]

for i, title in enumerate(headers):
    hdr_cells[i].text = title
    hdr_cells[i].width = widths[i]
    set_cell_background(hdr_cells[i], "1E3A8A") # Navy
    p = hdr_cells[i].paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.size = Pt(9.5)

driver_data = [
    ("VJMixConcretePlant121Driver.php", "V J Mix Concrete (Plant 121 - Schwing M1.5)", "1.25 m³", "8 Materials, 4 Sub-Runs"),
    ("PalaniyappaMci70Plant782Driver.php", "Palaniyappa Concrete (Plant 782 - MCI 70)", "0.50 m³", "6 Materials, 15 Sub-Runs"),
    ("PalaniyappaUdhayamPlant322Driver.php", "Palaniyappa / New Udhayam (Plant 322)", "2.50 m³", "6 Materials, 3 Sub-Runs"),
    ("SriGaneshaMci360PlantM1TDriver.php", "Sri Ganesha Readymix (Plant M1T-187)", "1.00 m³", "6 Materials, 5 Sub-Runs"),
    ("DynamicDbPlantDriver.php", "Generic Dynamic Database Fallback", "Dynamic", "Dynamic UI Mapping Template"),
]

for row_data in driver_data:
    row_cells = table.add_row().cells
    for j, val in enumerate(row_data):
        row_cells[j].text = val
        row_cells[j].width = widths[j]
        set_cell_margins(row_cells[j])
        p = row_cells[j].paragraphs[0]
        p.runs[0].font.size = Pt(8.5)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 4: Split-Screen User Experience
h1 = doc.add_heading("4. Split-Screen Verification & Approval Flow", level=1)
h1.paragraph_format.space_before = Pt(14)
h1.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph(
    "The UI incorporates a modern, full-width split-screen review modal (BatchSheetReview.vue):"
)

bp1 = doc.add_paragraph(style='List Bullet')
r = bp1.add_run("Left Half (Document Visualizer): ")
r.bold = True
bp1.add_run("Renders the exact uploaded PDF in an embedded high-resolution iframe, allowing operators to visually inspect and audit original physical figures.")

bp2 = doc.add_paragraph(style='List Bullet')
r = bp2.add_run("Right Half (Material & Header Matrix): ")
r.bold = True
bp2.add_run("Displays the auto-extracted headers and interactive Material Consumption table with live target vs actual calculations, variance kg/%, and tolerance badges.")

bp3 = doc.add_paragraph(style='List Bullet')
r = bp3.add_run("Load Tolerance Audit: ")
r.bold = True
bp3.add_run("Enforces the ±2.0% IS/ASTM batch tolerance standard. Highlights compliant batches in emerald green and flags excessive material deviations in amber.")

# Section 5: Database Entities & Schema
h1 = doc.add_heading("5. Database Architecture & Data Schema", level=1)
h1.paragraph_format.space_before = Pt(14)
h1.paragraph_format.space_after = Pt(6)

table2 = doc.add_table(rows=1, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.autofit = False

hdr_cells2 = table2.rows[0].cells
headers2 = ["Table Name", "Primary Purpose", "Key Columns"]
widths2 = [Inches(2.0), Inches(2.2), Inches(2.3)]

for i, title in enumerate(headers2):
    hdr_cells2[i].text = title
    hdr_cells2[i].width = widths2[i]
    set_cell_background(hdr_cells2[i], "334155") # Slate Dark
    p = hdr_cells2[i].paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.size = Pt(9.5)

schema_data = [
    ("mm_batch_sheet_uploads", "Staging and audit trail of raw uploads", "sha256_hash, status, parsed_json, normalized_json, confidence_score"),
    ("mm_batch_sheet_templates", "Saved plant layouts and material mappings", "plant_id, name, field_mapping, material_mapping, is_active"),
    ("mm_batches", "Core production load record", "plant_id, batch_no, batch_size, total_target_weight, total_actual_weight"),
    ("mm_batch_materials", "Physical consumption run rows", "batch_id, product_id, target_qty, actual_qty, deviation_quantity, runs"),
]

for row_data in schema_data:
    row_cells = table2.add_row().cells
    for j, val in enumerate(row_data):
        row_cells[j].text = val
        row_cells[j].width = widths2[j]
        set_cell_margins(row_cells[j])
        p = row_cells[j].paragraphs[0]
        p.runs[0].font.size = Pt(8.5)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 6: Verification & How to Operate
h1 = doc.add_heading("6. Operational Verification & Testing", level=1)
h1.paragraph_format.space_before = Pt(14)
h1.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph("The engine is fully verified and running live in development and production environments:")

bp1 = doc.add_paragraph(style='List Bullet')
r = bp1.add_run("Web Browser Flow: ")
r.bold = True
bp1.add_run("Open http://localhost:8000/batches -> Drop any of the 4 plant PDFs -> Split-screen modal opens -> Verify materials -> Click 'Approve & Save Batch'.")

bp2 = doc.add_paragraph(style='List Bullet')
r = bp2.add_run("Automated CLI Suite: ")
r.bold = True
bp2.add_run("Run 'php scratch/test_batch_sheet.php' to execute end-to-end regression tests against all plant samples.")

output_path = os.path.abspath("Batch_Sheet_Reconciliation_Engine_Documentation.docx")
doc.save(output_path)
print(f"Successfully generated DOCX at: {output_path}")
